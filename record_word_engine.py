# -*- coding: utf-8 -*-
from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Inches, RGBColor

from template_record_engine import TEMPLATE_DIR, fill_exact_template

BLACK = RGBColor(0, 0, 0)


def _signature_path(username: str):
    from lims_db import SIGNATURE_DIR, signature
    meta=signature(username) if username else None
    if not meta or not meta.get("image_file"):
        return None
    path=SIGNATURE_DIR/meta["image_file"]
    return path if path.exists() else None


def _put_signature(cell, username: str, date_text: str = ""):
    cell.text=""
    paragraph=cell.paragraphs[0]
    path=_signature_path(username)
    if path:
        paragraph.add_run().add_picture(str(path),width=Inches(.82))
    else:
        paragraph.add_run("【未配置签名图片】")
    if date_text:
        paragraph.add_run(f"  {str(date_text)[:10]}")


def _apply_record_signatures(doc, record):
    """Paste tester/reviewer handwriting into every actual signature field."""
    from lims_db import task
    task_row=task(record.get("task_no") or record.get("record_no")) or {}
    tester=record.get("owner") or task_row.get("assignee","")
    reviewer=task_row.get("reviewer","")
    tester_date=record.get("tester_signed_at") or record.get("updated_at","")
    reviewer_date=record.get("reviewer_signed_at") or ""
    tester_tokens=("检测人员","实验员")
    reviewer_tokens=("核验人员","复核人员","核验员","复核员")
    for table in doc.tables:
        if not table.rows:
            continue
        # Observer signature columns are true signatures, not typed name fields.
        headers=[cell.text.strip() for cell in table.rows[0].cells]
        observer_columns=[i for i,text in enumerate(headers) if "观察者签字" in text]
        for row in table.rows[1:]:
            for col in observer_columns:
                if col < len(row.cells) and row.cells[col].text.strip() not in ("","/","不适用"):
                    _put_signature(row.cells[col],tester,tester_date)
        for row_index,row in enumerate(table.rows):
            cells=[];seen=set()
            for index,cell in enumerate(row.cells):
                if cell._tc in seen:
                    continue
                seen.add(cell._tc);cells.append((index,cell,cell.text.strip()))
            row_text=" ".join(text for _index,_cell,text in cells)
            exact_role_row=any(
                text.strip(" ：:（）()") in tester_tokens+reviewer_tokens
                for _index,_cell,text in cells
            )
            signature_row=(
                any(token in row_text for token in ("签字","签名","日期","年__","年__月","/年/月"))
                or (row_index>0 and exact_role_row and len(cells)<=8)
            )
            if not signature_row:
                continue
            for position,(_index,_cell,label) in enumerate(cells):
                username="";signed_at=""
                if any(token in label for token in tester_tokens):
                    username,signed_at=tester,tester_date
                elif any(token in label for token in reviewer_tokens):
                    username,signed_at=reviewer,reviewer_date
                if not username or position+1>=len(cells):
                    continue
                target_position=position+1
                for later in range(position+1,len(cells)-1):
                    if cells[later][2].strip("：: ") in ("签字","签名"):
                        target_position=later+1
                        break
                _put_signature(cells[target_position][1],username,signed_at)
            # Combined confirmation cells in R004 contain both roles.
            for _index,cell,text in cells:
                if "确认人" in text and "复核" in text:
                    cell.text=""
                    p=cell.paragraphs[0]
                    p.add_run("确认：")
                    tester_path=_signature_path(tester)
                    if tester_path:p.add_run().add_picture(str(tester_path),width=Inches(.62))
                    p.add_run("  复核：")
                    reviewer_path=_signature_path(reviewer)
                    if reviewer_path:p.add_run().add_picture(str(reviewer_path),width=Inches(.62))


def _fallback(record):
    """Fallback is used only for a newly added experiment without a controlled template."""
    payload = record.get("payload", {})
    doc = Document()
    doc.add_heading(record.get("experiment") or "实验原始记录", 0)
    doc.add_paragraph("当前实验尚未配置受控原始记录模板。正式提交前应由管理员上传并发布模板版本。")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "字段"
    table.rows[0].cells[1].text = "记录值"
    for key, value in (payload.get("template_fields") or {}).items():
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(value)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = BLACK
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def export_record(record, template_name, changes):
    """Fill the controlled DOCX directly without adding, deleting or rebuilding its layout.

    The exported document contains only the original template content and values written into
    the template's existing cells. Equipment details and attachment indexes are not appended.
    """
    if not template_name or not (TEMPLATE_DIR / template_name).exists():
        return _fallback(record)

    changed_keys = set()
    if int(record.get("version", 1) or 1) > 1:
        for item in changes or []:
            field_name = str(item.get("field_name", ""))
            if item.get("action") == "字段修改" and field_name.startswith("template_fields."):
                changed_keys.add(field_name.split("template_fields.", 1)[1])

    values = (record.get("payload") or {}).get("template_fields") or {}
    doc = fill_exact_template(template_name, values, changed_keys)
    _apply_record_signatures(doc,record)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
