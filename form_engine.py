# -*- coding: utf-8 -*-
from __future__ import annotations
from io import BytesIO
from pathlib import Path
from typing import Any
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_ROW_HEIGHT_RULE
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from experiment_engine import result_summary
from report_rules import overall_conclusion, report_item

ROOT=Path(__file__).parent
TEMPLATE_DIR=ROOT/"templates"
SIGNATURE_DIR=ROOT/"data"/"signatures"
BLACK=RGBColor(0,0,0)


def _blacken(doc):
    for p in doc.paragraphs:
        for r in p.runs:r.font.color.rgb=BLACK
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:r.font.color.rgb=BLACK


def _save(doc):
    _blacken(doc);b=BytesIO();doc.save(b);b.seek(0);return b


def _setp(p,text,bold=False):
    p.clear();r=p.add_run(str(text));r.bold=bold;r.font.color.rgb=BLACK


def _prefix(p,prefix,value):
    if p.text.strip().startswith(prefix):_setp(p,prefix+("" if value is None else str(value)));return True
    return False


def _fill(table,data,header_rows=1):
    while len(table.rows)<header_rows+len(data):table.add_row()
    for i,vals in enumerate(data,start=header_rows):
        for j,v in enumerate(vals):
            if j<len(table.rows[i].cells):table.rows[i].cells[j].text="" if v is None else str(v)
    for i in range(header_rows+len(data),len(table.rows)):
        for cell in table.rows[i].cells:cell.text=""


def group_range(g):
    q=int(g.get("quantity") or 1)
    return f"{g['group_no']}-S01" if q==1 else f"{g['group_no']}-S01～{g['group_no']}-S{q:02d}"


def report_delivery_document(report_no: str, delivery_rows: list[dict[str, Any]]):
    """Fill the controlled report-delivery template supplied by the laboratory."""
    from lims_db import commission, commission_groups, one, report, report_actions
    d=Document(TEMPLATE_DIR/"FORM_REPORT_DELIVERY.docx")
    table=d.tables[0]
    report_row=report(report_no) or {}
    commission_row=commission(report_row.get("commission_no","")) or {}
    groups=commission_groups(report_row.get("commission_no","")) if report_row else []
    sample_text="；".join(
        f"{g.get('sample_name','')}/{g.get('model','')}" for g in groups
    )
    def checked(template,label):
        value=template.replace("☑","□")
        return value.replace("□"+label,"☑"+label,1)
    ordered=sorted(delivery_rows,key=lambda x:str(x.get("delivered_at","")))[:7]
    for offset in range(7):
        row=table.rows[offset+1]
        if offset>=len(ordered):
            for index in (1,2,4,7,10,11,13,14):
                _set_cell_existing(row.cells[index],"")
            continue
        item=ordered[offset]
        note=str(item.get("receipt_note",""))
        delivery_type="作废替换" if "作废替换" in note else ("更正" if "更正" in note else ("补发" if "补发" in note else "首次"))
        medium=str(commission_row.get("report_medium") or "电子")
        medium_label="纸质" if "纸" in medium else "电子"
        method=str(item.get("delivery_method",""))
        method_label="现场" if method in ("自取","现场领取","现场") else ("快递" if "快递" in method else ("邮件" if "邮件" in method else "系统"))
        values={
            0:offset+1,1:report_no,2:item.get("client_name") or commission_row.get("client_name",""),
            4:sample_text,5:checked("□首次 □补发 □更正 □作废替换",delivery_type),
            6:checked("□纸质 □电子\n份数：1",medium_label),
            7:item.get("delivered_at",""),9:checked("□现场 □快递 □邮件 □系统",method_label),
            10:f"{item.get('recipient','')} / {item.get('recipient_contact','')}",
            11:f"{item.get('receipt_status','')}；{note}".strip("；"),
            13:item.get("operator",""),14:note,
        }
        for index,value in values.items():
            _set_cell_existing(row.cells[index],value)
        _set_cell_signature(
            row.cells[13],item.get("operator"),item.get("delivered_at",""),width=.58,
        )
    latest=ordered[-1] if ordered else {}
    date_text=str(latest.get("delivered_at",""))[:10]
    manager_name=(one("SELECT display_name FROM users WHERE username=?",(latest.get("operator",""),)) or {}).get("display_name",latest.get("operator",""))
    approver_name=(one("SELECT display_name FROM users WHERE username=?",(report_row.get("approver",""),)) or {}).get("display_name",report_row.get("approver",""))
    _set_cell_existing(table.rows[8].cells[4],"☑审批签字完整  ☑报告编号一致  ☑页码完整  ☑专用章/电子章完整")
    _set_cell_existing(table.rows[8].cells[9],"☑附表齐全  ☑照片/附件齐全  ☑电子文件可正常打开")
    _set_cell_existing(table.rows[8].cells[16],"☑委托单位一致  ☑接收人信息正确  ☑交付方式符合约定")
    actions=[
        x for x in report_actions(report_no)
        if "作废" in str(x.get("action","")) or "更正" in str(x.get("action",""))
    ]
    latest_action=actions[-1] if actions else {}
    action_comment=str(latest_action.get("comment",""))
    changed=bool(latest_action)
    _set_cell_existing(table.rows[9].cells[4],f"{'□无  ☑有' if changed else '☑无  □有'}，说明：{action_comment if changed else ''}")
    _set_cell_existing(table.rows[9].cells[9],f"申请/批准记录编号：{report_no}-CHG-{latest_action.get('id','')}" if changed else "申请/批准记录编号：不适用")
    if "已收回" in action_comment:
        handling="☑收回  □作废  □无法收回已书面告知  □仅电子替换  □不适用"
    elif "无法收回" in action_comment:
        handling="□收回  ☑作废  ☑无法收回已书面告知  □仅电子替换  □不适用"
    elif "电子报告" in action_comment:
        handling="□收回  ☑作废  □无法收回已书面告知  ☑仅电子替换  □不适用"
    elif changed:
        handling="□收回  ☑作废  □无法收回已书面告知  □仅电子替换  □不适用"
    else:
        handling="□收回  □作废  □无法收回已书面告知  □仅电子替换  ☑不适用"
    _set_cell_existing(table.rows[9].cells[16],handling)
    _set_cell_signature(table.rows[10].cells[4],latest.get("operator"),date_text if ordered else "",width=.86)
    _set_cell_signature(table.rows[10].cells[9],report_row.get("approver"),report_row.get("approver_signed_at",""),width=.86)
    _set_cell_existing(table.rows[10].cells[16],f"电子路径：系统单据中心/{report_no}")
    _set_cell_existing(table.rows[11].cells[4],date_text)
    _set_cell_existing(table.rows[11].cells[9],str(report_row.get("approver_signed_at",""))[:10])
    return _save(d)


def modification_log_pdf(log_rows: list[dict[str, Any]], scope: str = "全部单据") -> BytesIO:
    """Create a Chinese-safe PDF without parsing log values as XML/HTML."""
    output=BytesIO()
    font_name="STSong-Light"
    for font_path in (
        ROOT/"assets"/"NotoSansCJK-Regular.ttc",
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf"),
        Path("/System/Library/Fonts/Supplemental/Songti.ttc"),
    ):
        if font_path.exists():
            try:
                font_name="BPLab-CJK"
                try:pdfmetrics.getFont(font_name)
                except Exception:pdfmetrics.registerFont(TTFont(font_name,str(font_path),subfontIndex=0))
                break
            except Exception:
                font_name="STSong-Light"
    if font_name=="STSong-Light":
        try:pdfmetrics.getFont(font_name)
        except Exception:pdfmetrics.registerFont(UnicodeCIDFont(font_name))
    page_size=landscape(A4)
    pdf=canvas.Canvas(output,pagesize=page_size)
    pdf.setTitle("修改记录日志")
    width,height=page_size
    margin=8*mm
    headers=["编号","单据/对象","修改位置","修改内容","修改前","修改后","原因","操作者/角色","服务器时间"]
    col_widths=[10,30,30,25,43,43,37,30,25]
    col_widths=[value*mm for value in col_widths]
    font_size=6.5
    line_height=8.5
    page_no=0

    def safe(value):
        text="" if value is None else str(value)
        return "".join(ch if ch in "\t\n\r" or ord(ch)>=32 else " " for ch in text)

    def wrap(text,max_width):
        lines=[]
        for source_line in safe(text).replace("\r","\n").split("\n"):
            if not source_line:
                lines.append("")
                continue
            current=""
            for char in source_line:
                if pdfmetrics.stringWidth(current+char,font_name,font_size)<=max_width-4:
                    current+=char
                else:
                    lines.append(current or char)
                    current="" if current else ""
                    if current=="" and pdfmetrics.stringWidth(char,font_name,font_size)<=max_width-4:
                        current=char
            lines.append(current)
        return lines or [""]

    def page_header():
        nonlocal page_no
        page_no+=1
        pdf.setFillColor(colors.HexColor("#12364A"))
        pdf.setFont(font_name,16)
        pdf.drawCentredString(width/2,height-margin-4,"修改记录日志")
        pdf.setFillColor(colors.black);pdf.setFont(font_name,8)
        pdf.drawString(margin,height-margin-20,f"范围：{safe(scope)}    记录数量：{len(log_rows)}")
        pdf.setFont(font_name,6.5);pdf.setFillColor(colors.HexColor("#475569"))
        pdf.drawString(margin,height-margin-32,"仅列示修改、作废、更正和照片替代；修改前后值及原因不可覆盖。")
        y=height-margin-44
        pdf.setFillColor(colors.HexColor("#176B87"))
        pdf.rect(margin,y-18,sum(col_widths),18,fill=1,stroke=0)
        x=margin
        pdf.setFillColor(colors.white);pdf.setFont(font_name,7)
        for label,col_width in zip(headers,col_widths):
            pdf.drawCentredString(x+col_width/2,y-12,label)
            x+=col_width
        return y-18

    def footer():
        pdf.setFillColor(colors.HexColor("#475569"));pdf.setFont(font_name,6.5)
        pdf.drawString(margin,6*mm,"系统完整审计链另行保存，可验证日志未被覆盖。")
        pdf.drawRightString(width-margin,6*mm,f"第 {page_no} 页")

    y=page_header()
    records=[]
    for item in log_rows:
        records.append([
            item.get("id",""),
            f"{item.get('entity_type','')} / {item.get('entity_id','')}",
            item.get("field_label") or item.get("field_name") or "单据级",
            item.get("action",""),item.get("old_value",""),item.get("new_value",""),
            item.get("reason",""),
            f"{item.get('actor_name') or item.get('actor','')} / {item.get('actor_role','')}",
            item.get("created_at",""),
        ])
    if not records:
        records=[["","暂无修改记录","","","","","","",""]]
    for row_number,values in enumerate(records):
        line_sets=[wrap(value,col_width) for value,col_width in zip(values,col_widths)]
        consumed=0;total=max(len(lines) for lines in line_sets)
        while consumed<total:
            available_lines=max(1,int((y-margin-7)/line_height))
            take=min(total-consumed,available_lines)
            row_height=max(18,take*line_height+6)
            if y-row_height<margin:
                footer();pdf.showPage();y=page_header()
                continue
            x=margin
            pdf.setFillColor(colors.HexColor("#F3F7F9") if row_number%2 else colors.white)
            pdf.rect(margin,y-row_height,sum(col_widths),row_height,fill=1,stroke=0)
            pdf.setStrokeColor(colors.HexColor("#94A3B8"));pdf.setLineWidth(0.35)
            pdf.setFont(font_name,font_size);pdf.setFillColor(colors.black)
            for lines,col_width in zip(line_sets,col_widths):
                pdf.rect(x,y-row_height,col_width,row_height,fill=0,stroke=1)
                for line_index,line in enumerate(lines[consumed:consumed+take]):
                    pdf.drawString(x+2,y-9-line_index*line_height,safe(line))
                x+=col_width
            y-=row_height
            consumed+=take
            if consumed<total:
                footer();pdf.showPage();y=page_header()
    footer();pdf.save();output.seek(0)
    return output


def hazardous_waste_document(item: dict[str, Any]):
    """Fill the controlled hazardous-waste template supplied by the laboratory."""
    from lims_db import one
    d=Document(TEMPLATE_DIR/"FORM_HAZARDOUS_WASTE.docx")
    task_nos=json.loads(item.get("task_nos") or "[]")
    occurred=str(item.get("occurred_at",""))
    date_text=occurred[:10]
    time_text=occurred[11:16]
    def option_line(text,label):
        value=text.replace("☑","□")
        token="□ "+label
        return value.replace(token,"☑ "+label,1) if token in value else value
    info=d.tables[1]
    _set_cell_existing(info.rows[0].cells[1],item.get("disposal_no",""))
    _set_cell_existing(info.rows[0].cells[3],date_text)
    _set_cell_existing(info.rows[1].cells[1],"实验室")
    _set_cell_existing(info.rows[2].cells[1],f"{item.get('commission_no','')} / {'、'.join(task_nos)}")
    _set_cell_existing(info.rows[2].cells[3],"、".join(task_nos))
    _set_cell_existing(info.rows[3].cells[1],item.get("handler",""))
    _set_cell_existing(info.rows[3].cells[3],item.get("container_no",""))
    source=d.tables[2].cell(0,0)
    source_text=source.text
    source_text=option_line(source_text,"检测实验")
    waste_type=str(item.get("waste_type",""))
    mapping={"实验废液":"实验废液","废弃样品":"废弃/破坏性样品","沾染耗材":"受污染耗材/容器"}
    source_text=option_line(source_text,mapping.get(waste_type,"其他"))
    source_text=source_text.replace(
        "产生情况及成分说明：______________________________________________________________",
        f"产生情况及成分说明：{item.get('waste_name','')}；{item.get('note','')}",
    )
    _set_cell_existing(source,source_text)
    detail=d.tables[3]
    for ri in range(1,7):
        for ci in range(1,9):
            _set_cell_existing(detail.rows[ri].cells[ci],"")
    _set_cell_existing(detail.rows[1].cells[1],item.get("disposal_no",""))
    _set_cell_existing(detail.rows[1].cells[2],"、".join(task_nos))
    _set_cell_existing(detail.rows[1].cells[3],item.get("waste_name",""))
    _set_cell_existing(detail.rows[1].cells[4],"液体" if "液" in waste_type else "固体")
    _set_cell_existing(detail.rows[1].cells[5],f"{item.get('quantity','')} {item.get('unit','')}")
    _set_cell_existing(detail.rows[1].cells[6],item.get("container_no",""))
    _set_cell_existing(detail.rows[1].cells[7],item.get("container_no",""))
    _set_cell_existing(detail.rows[1].cells[8],item.get("note",""))
    classification=d.tables[4]
    _set_cell_existing(classification.rows[0].cells[1],option_line(classification.rows[0].cells[1].text,"危险废物"))
    _set_cell_existing(classification.rows[1].cells[1],option_line(classification.rows[1].cells[1].text,"液体" if "液" in waste_type else "固体"))
    _set_cell_existing(classification.rows[2].cells[1],option_line(classification.rows[2].cells[1].text,mapping.get(waste_type,"废弃样品/制样残余物")))
    _set_cell_existing(classification.rows[6].cells[1],item.get("hazard_category",""))
    _set_signature_paragraph(
        classification.rows[7].cells[1].paragraphs[0],
        item.get("handler"),date_text,label="实验员：",width=.82,
    )
    storage=d.tables[5]
    _set_cell_existing(storage.rows[0].cells[3],item.get("container_no",""))
    _set_cell_existing(storage.rows[1].cells[1],item.get("container_no",""))
    disposal=d.tables[7]
    _set_cell_existing(disposal.rows[1].cells[1],f"{date_text} {time_text}")
    _set_cell_existing(disposal.rows[2].cells[1],item.get("container_no",""))
    _set_cell_existing(disposal.rows[3].cells[1],item.get("disposal_method",""))
    _set_cell_existing(disposal.rows[4].cells[1],f"{item.get('quantity','')} {item.get('unit','')}")
    _set_cell_existing(disposal.rows[5].cells[1],f"☑ 已完成分类包装  状态：{item.get('status','')}")
    _set_cell_existing(disposal.rows[6].cells[1],item.get("note",""))
    _set_signature_paragraph(
        disposal.rows[7].cells[1].paragraphs[0],
        item.get("handler"),date_text,label="执行人员：",width=.82,
    )
    signatures=d.tables[8]
    first_task=task_nos[0] if task_nos else item.get("task_no","")
    task_roles=one(
        "SELECT reviewer FROM tasks WHERE task_no=?",(first_task,),
    ) or {}
    approver=one(
        "SELECT username FROM users WHERE role='管理员' AND enabled=1 ORDER BY username LIMIT 1"
    ) or {}
    _set_cell_signature(signatures.rows[1].cells[1],item.get("handler"),date_text,width=.82)
    _set_cell_existing(signatures.rows[1].cells[2],date_text)
    _set_cell_signature(signatures.rows[2].cells[1],task_roles.get("reviewer"),date_text,width=.82)
    _set_cell_existing(signatures.rows[2].cells[2],date_text)
    _set_cell_signature(signatures.rows[3].cells[1],approver.get("username"),date_text,width=.82)
    _set_cell_existing(signatures.rows[3].cells[2],date_text)
    return _save(d)


def _objection_form(
    title: str, form_no: str,
    sections: list[tuple[str, list[tuple[str, Any]]]],
    signers: list[tuple[str, str, str]] | None = None,
):
    d=Document()
    section=d.sections[0]
    section.top_margin=Inches(0.55);section.bottom_margin=Inches(0.55)
    title_p=d.add_paragraph();title_p.alignment=1
    run=title_p.add_run(title);run.bold=True;run.font.size=Pt(17)
    sub=d.add_paragraph(f"表单编号：{form_no}");sub.alignment=2
    for section_title, fields in sections:
        heading=d.add_paragraph();r=heading.add_run(section_title);r.bold=True;r.font.size=Pt(11)
        table=d.add_table(rows=0,cols=2);table.style="Table Grid"
        for label,value in fields:
            cells=table.add_row().cells
            cells[0].width=Inches(1.55);cells[1].width=Inches(5.65)
            cells[0].text=str(label);cells[1].text="" if value is None else str(value)
            for rr in cells[0].paragraphs[0].runs:rr.bold=True
    if signers:
        heading=d.add_paragraph();r=heading.add_run("电子签名");r.bold=True;r.font.size=Pt(11)
        table=d.add_table(rows=0,cols=3);table.style="Table Grid"
        for label,person,date_text in signers:
            cells=table.add_row().cells
            cells[0].text=label
            _set_cell_signature(cells[1],person,date_text,width=.88)
            cells[2].text=str(date_text or "")[:10]
    d.add_paragraph("系统留痕：本表内容与异议流程记录、追溯资料及修改日志关联保存。")
    return _save(d)


def objection_application_document(obj: dict[str, Any], report_row: dict[str, Any], commission_row: dict[str, Any]):
    return _objection_form("客户异议申请表",obj.get("objection_no",""),[
        ("一、申请信息",[
            ("异议编号",obj.get("objection_no","")),("关联报告编号",obj.get("report_no","")),
            ("委托编号",obj.get("commission_no","")),("委托单位",obj.get("client_name","")),
            ("联系人",obj.get("contact","")),("联系电话",commission_row.get("phone","")),
            ("申请日期",obj.get("submitted_at","")),("受理渠道",obj.get("application_channel","")),
            ("争议项目",obj.get("disputed_items","")),("涉及样品",obj.get("involved_samples","")),
        ]),
        ("二、客户异议内容",[
            ("异议描述",obj.get("description","")),("随附材料/证据",obj.get("evidence_note","")),
        ]),
        ("三、受理信息",[
            ("登记人",obj.get("registered_by","")),("质量调查人",obj.get("quality_inspector","")),
            ("当前状态",obj.get("status","")),("登记时间",obj.get("created_at","")),
        ]),
    ],[
        ("样品管理员受理签名",obj.get("registered_by",""),obj.get("created_at","")),
        ("质量负责人调查签名",obj.get("quality_inspector",""),obj.get("investigated_at","")),
    ])


def objection_response_document(obj: dict[str, Any], report_row: dict[str, Any], commission_row: dict[str, Any]):
    return _objection_form("客户异议回复单",f"{obj.get('objection_no','')}-R",[
        ("一、异议基本信息",[
            ("异议编号",obj.get("objection_no","")),("关联报告编号",obj.get("report_no","")),
            ("委托单位",obj.get("client_name","")),("联系人",obj.get("contact","")),
        ]),
        ("二、调查与处理",[
            ("责任判定",obj.get("pathway","")),("调查过程",obj.get("investigation","")),
            ("调查结论",obj.get("trace_conclusion","")),("影响范围",obj.get("impact_scope","")),
            ("客户重测决定",obj.get("customer_retest_decision","")),
            ("重测任务/替代报告",f"{obj.get('retest_task_no','')} / {obj.get('replacement_report_no','')}"),
        ]),
        ("三、正式回复",[
            ("回复内容",obj.get("response_text","")),("回复方式",obj.get("response_method","")),
            ("样品管理员",obj.get("sent_by","")),("发送时间",obj.get("sent_at","")),
            ("接收凭证/备注",obj.get("response_receipt","")),
        ]),
    ],[
        ("质量负责人调查签名",obj.get("quality_inspector",""),obj.get("investigated_at","")),
        ("样品管理员回复签名",obj.get("sent_by","") or obj.get("registered_by",""),obj.get("sent_at","")),
    ])


def commission_document(c,groups,tests,receiver_name):
    d=Document(TEMPLATE_DIR/"FORM_COMMISSION.docx")
    methods=json.loads(c.get("method_choices") or "[]")
    options=["YY/T 1936","YY 0300","YY 0621.1","YY 0621.2","YY/T 1702","GB 17168","GB/T 4340.1","GB/T 3851","GB/T 18876.1","YY/T 1937","YY 0270.1","T/GDMDMA 0003","YY 0710"]
    method_line1="  ".join(("☑" if x in methods else "□")+x for x in options[:7])
    method_line2="  ".join(("☑" if x in methods else "□")+x for x in options[7:])
    bad=[g for g in groups if g.get("condition")!="完好"]
    bad_note="；".join(f"{g['group_no']}:{g.get('condition_note','')}" for g in bad)
    for p in d.paragraphs:
        _prefix(p,"委托方名称：",c.get("client_name",""));_prefix(p,"委托方地址：",c.get("client_address",""))
        if p.text.strip().startswith("联系人："):_setp(p,f"联系人：{c.get('contact','')}    联系电话：{c.get('phone','')}    委托日期：{c.get('commission_date','')}")
        if "样品外观检查良好" in p.text:_setp(p,f"{'□' if bad else '☑'}样品外观检查良好； {'☑' if bad else '□'}样品外观异常。异常情况说明：{bad_note}")
        if p.text.strip().startswith("检测方法："):_setp(p,"检测方法：")
        elif "YY/T 1936" in p.text:_setp(p,method_line1)
        elif "GB/T 3851" in p.text:_setp(p,method_line2)
        if p.text.strip().startswith("1、标普检测资源不满足时"):
            yes=c.get("subcontract_allowed")=="是";_setp(p,f"1、标普检测资源不满足时，是否允许分包？ {'☑是  □否' if yes else '□是  ☑否'}")
        if p.text.strip().startswith("2、样品、相关资料是否保密"):_setp(p,"2、样品、相关资料是否保密？ □是  ☑无要求")
        if p.text.strip().startswith("报告载体："):_setp(p,f"报告载体：{c.get('report_medium','')}    符合性判定：{c.get('conformity_judgment','')}")
        if p.text.strip().startswith("考虑不确定度："):_setp(p,f"考虑不确定度：{c.get('uncertainty','')}    递送方式：{c.get('delivery_method','')}")
        if "CNAS" in p.text and "章" in p.text:_setp(p,f"加盖 CNAS 章：{c.get('cnas_mark','')}")
        if p.text.strip().startswith("检测能力："):_setp(p,f"检测能力：{c.get('capability','')}")
    tm={}
    for t in tests:tm.setdefault(t["group_no"],[]).append(t["experiment"])
    data=[]
    for i,g in enumerate(groups,1):
        prod=c.get("production_org_name","")+("（受委托生产企业）" if c.get("production_relation")=="受委托生产企业" else "")
        data.append([i,f"{g['sample_name']}（{g['model']}）",group_range(g),prod,"、".join(tm.get(g["group_no"],[])),g.get("quantity",1),g.get("notes","") or g.get("condition_note","")])
    _fill(d.tables[0],data)
    table=d.tables[0]
    # Long model/specification values must wrap and remain fully visible.
    for row in table.rows[1:]:
        row.height=Inches(0.46)
        row.height_rule=WD_ROW_HEIGHT_RULE.AT_LEAST
        for paragraph in row.cells[3].paragraphs:
            paragraph.paragraph_format.keep_together=True
            for run in paragraph.runs:
                run.font.size=Pt(8)
    confirmation=d.tables[1].rows[0].cells[1]
    if len(confirmation.paragraphs)>=4:
        _set_signature_paragraph(
            confirmation.paragraphs[3],receiver_name,c.get("commission_date",""),width=1.05,
        )
    return _save(d)


def sample_register_document(c,groups,samples,tests,receiver_name):
    """Generate DLBP-CX-P10-R01 sample registration form.

    Column order follows the controlled template exactly:
    laboratory sample number, commissioning unit, sample name, model/specification,
    production unit, sample number/batch, inspection items, quantity, receiver,
    date and remarks.
    """
    d=Document(TEMPLATE_DIR/"FORM_SAMPLE_REGISTER.docx")
    gm={g["group_no"]:g for g in groups}
    tm={}
    for t in tests:
        tm.setdefault(t["group_no"],[]).append(t["experiment"])

    production_unit=c.get("production_org_name","")
    if c.get("production_relation")=="受委托生产企业" and production_unit:
        production_unit += "（受委托生产企业）"

    data=[]
    for s in samples:
        g=gm[s["group_no"]]
        remarks=s.get("condition_note","") or g.get("notes","") or ""
        data.append([
            s["sample_no"],
            c.get("client_name",""),
            s["sample_name"],
            s["model"],
            production_unit,
            g.get("product_no",""),
            "、".join(tm.get(s["group_no"],[])),
            1,
            receiver_name,
            c.get("commission_date",""),
            remarks,
        ])
    _fill(d.tables[0],data)
    for row in d.tables[0].rows[1:1+len(data)]:
        _set_cell_signature(
            row.cells[8],receiver_name,c.get("commission_date",""),width=.72,
        )
    return _save(d)


def loan_return_document(loans,user_names):
    d=Document(TEMPLATE_DIR/"FORM_SAMPLE_LOAN_RETURN.docx");data=[]
    for i,x in enumerate(loans,1):
        purpose=x.get("purpose") or "、".join(json.loads(x.get("experiments") or "[]"))
        data.append([i,x["sample_no"],user_names.get(x.get("borrower"),x.get("borrower","")),x.get("borrowed_at",""),purpose,x.get("returned_at",""),user_names.get(x.get("returned_by"),x.get("returned_by","")),x.get("return_note","") or x.get("issue_note","")])
    _fill(d.tables[0],data)
    for index,item in enumerate(loans,1):
        row=d.tables[0].rows[index]
        _set_cell_signature(
            row.cells[2],item.get("borrower"),item.get("borrowed_at",""),width=.62,
        )
        if item.get("returned_by"):
            _set_cell_signature(
                row.cells[6],item.get("returned_by"),item.get("returned_at",""),width=.62,
            )
    return _save(d)


def _sig(meta):
    if not meta or not meta.get("image_file"):return None
    p=SIGNATURE_DIR/meta["image_file"];return p if p.exists() else None


def _signature_for_person(person):
    """Resolve either a username or a display name to its controlled image."""
    if not person:
        return None
    from lims_db import one, signature
    user=one(
        "SELECT username FROM users WHERE username=? OR display_name=? LIMIT 1",
        (str(person),str(person)),
    )
    return _sig(signature(user["username"])) if user else None


def _set_signature_paragraph(paragraph, person, date_text="", label="", width=0.92):
    """Write a picture-only signature; no typed personal name is retained."""
    paragraph.clear()
    if label:
        paragraph.add_run(label)
    path=_signature_for_person(person)
    if path:
        try:
            paragraph.add_run().add_picture(str(path),width=Inches(width))
        except Exception:
            paragraph.add_run("【签名图片读取失败】")
    else:
        paragraph.add_run("【未配置签名图片】")
    if date_text:
        paragraph.add_run(f"  {str(date_text)[:10]}")


def _set_cell_signature(cell, person, date_text="", width=0.92):
    paragraphs=list(cell.paragraphs)
    paragraph=paragraphs[0] if paragraphs else cell.add_paragraph()
    _set_signature_paragraph(paragraph,person,date_text,width=width)
    for extra in paragraphs[1:]:
        _set_existing_text(extra,"")


def _set_existing_text(paragraph, text):
    """Replace text without removing the template paragraph/run elements."""
    runs=list(paragraph.runs)
    if not runs:
        runs=[paragraph.add_run("")]
    runs[0].text=str(text or "")
    runs[0].font.color.rgb=BLACK
    for run in runs[1:]:
        run.text=""
        run.font.color.rgb=BLACK


def _prefix_existing(paragraph, prefix, value):
    if paragraph.text.strip().startswith(prefix):
        _set_existing_text(paragraph,prefix+("" if value is None else str(value)))
        return True
    return False


def _set_cell_existing(cell, value):
    paragraphs=list(cell.paragraphs)
    paragraph=paragraphs[0] if paragraphs else cell.add_paragraph()
    _set_existing_text(paragraph,"" if value is None else str(value))
    for extra in paragraphs[1:]:
        _set_existing_text(extra,"")


def _fill_existing_rows(table, data, header_rows=1):
    """Fill only rows already present in the controlled report mother."""
    capacity=max(0,len(table.rows)-header_rows)
    for offset in range(capacity):
        values=data[offset] if offset<len(data) else []
        row=table.rows[header_rows+offset]
        for col,cell in enumerate(row.cells):
            _set_cell_existing(cell,values[col] if col<len(values) else "")


def _date_range(values):
    clean=sorted(dict.fromkeys(str(x)[:10] for x in values if x))
    if not clean:return ""
    return clean[0] if len(clean)==1 else f"{clean[0]}至{clean[-1]}"


def _range_text(values, suffix=""):
    nums=[]
    for value in values:
        try:nums.append(float(value))
        except Exception:pass
    if not nums:return ""
    lo,hi=min(nums),max(nums)
    def fmt(x):return f"{x:.3f}".rstrip("0").rstrip(".")
    return f"{fmt(lo)}{suffix}" if lo==hi else f"{fmt(lo)}～{fmt(hi)}{suffix}"


def report_document(c,groups,samples,tasks,records,report,user_names,signatures):
    d=Document(TEMPLATE_DIR/"FORM_REPORT.docx")
    names="、".join(dict.fromkeys(g["sample_name"] for g in groups));models="、".join(dict.fromkeys(g["model"] for g in groups));ranges="；".join(group_range(g) for g in groups);products="、".join(dict.fromkeys(g.get("product_no","") for g in groups if g.get("product_no")));prods=c.get("production_org_name","")+("（受委托生产企业）" if c.get("production_relation")=="受委托生产企业" else "");conds="；".join(f"{g['group_no']}:{g['condition']}" for g in groups)
    test_dates=[];report_items=[];equipment_map={};environment_by_location={}
    for t in tasks:
        rec=records.get(t["task_no"])
        if not rec:continue
        payload=rec.get("payload") or {}
        business=payload.get("business_record") or {}
        params=business.get("parameters") or payload.get("parameters") or {}
        rows=business.get("rows") or payload.get("data") or []
        common=payload.get("common") or {}
        snapshot=payload.get("configuration_snapshot") or {}
        kind=t.get("kind") or snapshot.get("kind") or "generic"
        item=report_item(kind,rows)
        item.update({
            "task":t,
            "standard":t.get("standard") or snapshot.get("standard") or t.get("method_code",""),
            "deviation":business.get("deviation") or payload.get("deviation") or "无",
        })
        report_items.append(item)
        test_dates.append(params.get("test_date") or common.get("test_date"))
        location=(
            t.get("detection_location")
            or snapshot.get("selected_detection_location")
            or params.get("detection_location")
            or snapshot.get("default_location")
            or "未记录地点"
        )
        env=environment_by_location.setdefault(location,{"temperature":[],"humidity":[],"other":[]})
        env["temperature"].extend([params.get("temperature_before"),params.get("temperature_after")])
        env["humidity"].extend([params.get("humidity_before"),params.get("humidity_after")])
        if params.get("environment_interference"):
            env["other"].append(str(params.get("environment_interference")))
        check_map={x.get("management_no"):x for x in business.get("equipment_checks") or []}
        for eq in snapshot.get("equipment") or []:
            no=eq.get("management_no","")
            value=dict(eq);value["usage_status"]=(check_map.get(no) or {}).get("status","正常")
            equipment_map.setdefault(no or eq.get("equipment_name",""),value)

    auto_statement="；".join(f"{g['group_no']}：接收状态{g.get('condition','完好')}，共{g.get('quantity',1)}件" for g in groups)
    auto_conclusion=overall_conclusion(report_items)
    for p in d.paragraphs:
        if p.text.strip().startswith("生 产 单 位："):
            _set_existing_text(p,f"生 产 单 位：{prods}    接 收 日 期：{c.get('commission_date','')}")
            continue
        _prefix_existing(p,"报告编号：",report.get("report_no",""));_prefix_existing(p,"委 托 单 位：",c.get("client_name",""));_prefix_existing(p,"地       址：",c.get("client_address",""));_prefix_existing(p,"样 品 名 称：",names);_prefix_existing(p,"型 号/规 格：",models);_prefix_existing(p,"样 品 编 号：",ranges);_prefix_existing(p,"产品编号/批号：",products);_prefix_existing(p,"生 产 单 位：",prods);_prefix_existing(p,"接 收 日 期：",c.get("commission_date",""));_prefix_existing(p,"接 收 状 态：",conds);_prefix_existing(p,"检 验 类 别：",report.get("report_category") or "委托检验");_prefix_existing(p,"报告发布日期：",report.get("publish_date") or "")
        _prefix_existing(p,"检验日期 ：",_date_range(test_dates))
        if p.text.strip().startswith("需说明的情况:"):_set_existing_text(p,"需说明的情况:"+(report.get("notes") or "无"))
        if p.text.strip().startswith("样品情况说明："):_set_existing_text(p,"样品情况说明："+(report.get("sample_statement") or auto_statement))
        if p.text.strip().startswith("检验结论："):_set_existing_text(p,"检验结论："+(report.get("conclusion") or auto_conclusion))

    # Insert conclusion-driving evidence photos. REPORT_PHOTO is a fallback only.
    from lims_db import list_attachments, attachment_file
    from constants import REPORT_DECISIVE_PHOTO_CODES
    report_photos=[]
    for task_item in tasks:
        codes=REPORT_DECISIVE_PHOTO_CODES.get(task_item.get("experiment",""),[])
        all_candidates=[
            item for item in list_attachments(task_no=task_item["task_no"])
            if item.get("capture_source")=="live_camera"
            and item.get("evidence_status")=="有效"
            and not bool(item.get("is_original"))
        ]
        candidates=[item for item in all_candidates if item.get("checkpoint_code") in codes]
        if not candidates:
            candidates=[item for item in all_candidates if item.get("checkpoint_code")=="REPORT_PHOTO"]
        # Keep every sample-level decisive result, but de-duplicate task-level checkpoints.
        seen=set()
        for photo in sorted(candidates,key=lambda x:(codes.index(x.get("checkpoint_code")) if x.get("checkpoint_code") in codes else 999,x.get("sample_no",""),x.get("server_captured_at",""))):
            identity=(photo.get("checkpoint_code"),photo.get("sample_no") or "TASK")
            if identity in seen:
                continue
            seen.add(identity)
            report_photos.append((task_item,photo))
    for paragraph in d.paragraphs:
        if paragraph.text.strip().startswith("照片和说明："):
            _set_existing_text(paragraph,"照片和说明：")
            for task_item,photo in report_photos:
                path=attachment_file(photo)
                if not path.exists():
                    continue
                paragraph.add_run("\n")
                try:
                    paragraph.add_run().add_picture(str(path),width=Inches(2.6))
                    sample_text=photo.get("sample_no") or "任务整体"
                    record_item=next((x for x in report_items if x.get("task",{}).get("task_no")==task_item.get("task_no")),None) or {}
                    paragraph.add_run(
                        f"\n{task_item.get('experiment','')}｜样品：{sample_text}"
                        f"\n证据：{photo.get('checkpoint_label','')}"
                        f"\n时间：{str(photo.get('server_captured_at','')).replace('T',' ')}｜"
                        f"结论：{record_item.get('conclusion','')}"
                    )
                except Exception:
                    paragraph.add_run(f"\n照片文件：{photo.get('original_name','')}")
            if not report_photos:
                paragraph.add_run("未留档")
            break

    equipment_rows=[]
    equipment=list(equipment_map.values())
    for item in equipment[:5]:
        equipment_rows.append([
            f"{item.get('equipment_name','')}（{item.get('management_no','')}）",
            item.get("model",""),
            item.get("measuring_range",""),
            item.get("calibration_certificate") or "台账未配置",
            item.get("traceability_agency") or "台账未配置",
            item.get("calibration_due") or (f"台账校准时间：{item.get('calibration_time')}" if item.get("calibration_time") else "台账未配置"),
        ])
    environment_rows=[
        [location,_range_text(data["temperature"]," ℃"),_range_text(data["humidity"]," %RH"),"、".join(dict.fromkeys(data["other"])) or "无"]
        for location,data in environment_by_location.items()
    ][:3]

    grouped={}
    for item in report_items:
        group_no=item["task"].get("group_no","")
        grouped.setdefault(group_no,[]).append(item)
    result_rows=[]
    for index,(group_no,items) in enumerate(grouped.items(),1):
        experiments="；".join(f"{x['task'].get('experiment','')}（{x['standard']}）" for x in items)
        requirements="；".join(dict.fromkeys(x["requirement"] for x in items))
        actual="；".join(x["result"] for x in items)
        conclusions=[x["conclusion"] for x in items]
        conclusion="不符合" if any(x in ("不符合","不合格") for x in conclusions) else ("符合" if all(x in ("符合","合格") for x in conclusions) else "仅描述结果")
        notes="；".join(dict.fromkeys(x["deviation"] for x in items if x["deviation"] not in ("","无"))) or "无"
        result_rows.append([index,f"{group_no} {experiments}",requirements,actual,conclusion,notes])
    if len(result_rows)>3:
        overflow=result_rows[2:]
        result_rows=result_rows[:2]+[[
            3,
            "；".join(str(x[1]) for x in overflow),
            "；".join(str(x[2]) for x in overflow),
            "；".join(str(x[3]) for x in overflow),
            "不符合" if any(x[4]=="不符合" for x in overflow) else "符合",
            "；".join(str(x[5]) for x in overflow),
        ]]
    if len(d.tables)>0:_fill_existing_rows(d.tables[0],equipment_rows)
    if len(d.tables)>1:_fill_existing_rows(d.tables[1],environment_rows)
    if len(d.tables)>2:
        standards="；".join(dict.fromkeys(item["standard"] for item in report_items if item["standard"]))
        _set_cell_existing(d.tables[2].rows[0].cells[2],standards)
        _fill_existing_rows(d.tables[2],result_rows,2)
    source_records=[item for item in records.values() if item]
    record_tester_signed=max(
        (str(item.get("tester_signed_at") or "") for item in source_records),default="",
    )
    record_reviewer_signed=max(
        (str(item.get("reviewer_signed_at") or "") for item in source_records),default="",
    )
    signature_rows=[
        ("批 准 人",report.get("approver"),report.get("approver_signed_at")),
        # 报告首页的“核验员”就是原始记录复核员。
        ("核 验 员",report.get("verifier"),report.get("verifier_signed_at") or record_reviewer_signed),
        # 报告首页的“检测员”就是执行该实验的实验员。
        ("检 测 员",report.get("tester"),report.get("tester_signed_at") or record_tester_signed),
    ]
    for label,u,signed_value in signature_rows:
        for p in d.paragraphs:
            if p.text.strip().startswith(label):
                signed_date=str(signed_value or "")[:10]
                if signed_value and u:
                    _set_signature_paragraph(
                        p,u,signed_date,label=f"{label}    ",width=.92,
                    )
                else:
                    _set_existing_text(p,f"{label}    【待签名】")
                break
    return _save(d)
